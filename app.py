import os,io,re,streamlit as st
from typing import TypedDict
from langchain_anthropic import ChatAnthropic
from langchain_community.tools.tavily_search import TavilySearchResults
from langgraph.graph import StateGraph,END
from docx import Document
from docx.oxml.shared import OxmlElement,qn
from docx.opc.constants import RELATIONSHIP_TYPE
from pypdf import PdfReader
import streamlit.components.v1 as components

os.environ["TAVILY_API_KEY"]=st.secrets["TAVILY_API_KEY"]
os.environ["ANTHROPIC_API_KEY"]=st.secrets["ANTHROPIC_API_KEY"]
st.set_page_config(page_title="Enterprise Strategic Intelligence",layout="wide")
st.markdown("""
<link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.2/dist/css/bootstrap.min.css" rel="stylesheet">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">

<style>

section[data-testid="stSidebar"]{display:none;}
div[data-testid="collapsedControl"]{display:none;}
header{display:none;}
footer{display:none;}

html,body,.stApp{
 font-family:'Inter',sans-serif;
 background:#020617;
 color:#e2e8f0;
 overflow-x:hidden;
}

/* 3D background canvas */
#bg-canvas{
 position:fixed;
 top:0;
 left:0;
 width:100%;
 height:100%;
 z-index:-1;
}

/* Navbar */

.navbar-custom{
 display:flex;
 justify-content:space-between;
 align-items:center;
 padding:20px 60px;
 background:rgba(255,255,255,0.03);
 backdrop-filter:blur(20px);
 border-bottom:1px solid rgba(255,255,255,0.05);
}

.logo{
 font-weight:600;
 font-size:22px;
 background:linear-gradient(90deg,#38bdf8,#6366f1);
 -webkit-background-clip:text;
 -webkit-text-fill-color:transparent;
}

/* Hero */

.hero{
 text-align:center;
 margin-top:80px;
 margin-bottom:40px;
}

.hero-title{
 font-size:52px;
 font-weight:700;
 background:linear-gradient(90deg,#38bdf8,#818cf8);
 -webkit-background-clip:text;
 -webkit-text-fill-color:transparent;
}

.hero-sub{
 color:#94a3b8;
 margin-top:10px;
}

/* Glass cards */

.glass-card{
 max-width:1100px;
 margin-left:auto;
 margin-right:auto;
 background:rgba(255,255,255,0.05);

 backdrop-filter:blur(20px);

 border:1px solid rgba(255,255,255,0.08);

 border-radius:16px;

 padding:30px;

 margin-top:20px;

 transition:all 0.3s ease;

 animation:fadeUp 0.8s ease;

}

.glass-card:hover{

 transform:translateY(-5px);

 box-shadow:0 20px 40px rgba(0,0,0,0.6);

}

/* Inputs */

.stTextInput input{

 background:rgba(255,255,255,0.05);

 border:none;

 border-radius:10px;

 color:white;

 padding:12px;

}

.stFileUploader{

 background:rgba(255,255,255,0.05);

 border-radius:10px;

 padding:12px;

}

/* Button */

.stButton>button{

 background:linear-gradient(90deg,#2563eb,#7c3aed);

 border:none;

 border-radius:12px;

 padding:14px 32px;

 color:white;

 font-weight:600;

 transition:0.3s;

}

.stButton>button:hover{

 transform:scale(1.05);

 box-shadow:0 10px 30px rgba(99,102,241,0.5);

}

/* Animation */

@keyframes fadeUp{

 from{

 opacity:0;

 transform:translateY(40px);

 }

 to{

 opacity:1;

 transform:translateY(0);

 }

}

</style>

<div id="bg-canvas"></div>

<div class="navbar-custom">

<div class="logo">Report Generator </div>

<div>Deep Research</div>

</div>

<div class="hero">

<div class="hero-title">Enterprise Deep Research Intelligence Platform</div>

<div class="hero-sub">

AI-powered consulting-grade research and analysis

</div>

</div>

<script src="https://cdnjs.cloudflare.com/ajax/libs/three.js/r134/three.min.js"></script>

<script>

const scene=new THREE.Scene();

const camera=new THREE.PerspectiveCamera(75,window.innerWidth/window.innerHeight,0.1,1000);

const renderer=new THREE.WebGLRenderer({alpha:true});

renderer.setSize(window.innerWidth,window.innerHeight);

document.getElementById("bg-canvas").appendChild(renderer.domElement);

const geometry=new THREE.IcosahedronGeometry(2,1);

const material=new THREE.MeshStandardMaterial({

 color:0x6366f1,

 wireframe:true

});

const sphere=new THREE.Mesh(geometry,material);

scene.add(sphere);

const light=new THREE.PointLight(0x38bdf8,1);

light.position.set(5,5,5);

scene.add(light);

camera.position.z=5;

function animate(){

 requestAnimationFrame(animate);

 sphere.rotation.x+=0.003;

 sphere.rotation.y+=0.003;

 renderer.render(scene,camera);

}

animate();

</script>

""",unsafe_allow_html=True)

class OverallState(TypedDict):
 target_company:str
 pdf_context:str
 remaining_sections:list
 completed_research:list
 all_urls:list
 final_report:str
def get_llm():
 return ChatAnthropic(model="claude-3-haiku-20240307",temperature=0)
PROMPT_SOP={
"Section 1: Account Business Overview":"""Act as a senior banking analyst and strategy consultant.Combine factual accuracy with consulting insight.Use EXACT structure:
Section 1: Account Business Overview
1.1 Who they are
● **Who they are:** Description including holding company,ticker [ref](URL)
● **Founded:** Founding year [ref](URL)
● **Footprint:** Branch and geography [ref](URL)
1.2 Headquarters and regulator signals
● **HQ:** Location [ref](URL)
● **Primary regulator:** Name [ref](URL)
1.3 Scale and performance
● **Assets:** Value [ref](URL)
● **Deposits:** Value [ref](URL)
● **Net income:** Value [ref](URL)
1.4 Business model
● **Revenue drivers:** Lending,wealth,etc [ref](URL)
**Why this matters for Speridian**
Explain strategic implications and modernization opportunities [ref](URL)
STRICT RULES:
Every fact MUST include clickable [ref](URL)
Use only verified facts
DO NOT invent data""",
"Section 2: Account Key Business Initiatives":"""Act as strategy consultant.Use EXACT structure:
Section 2: Account Key Business Initiatives
2.1 Executive synthesis
Explain strategic priorities [ref](URL)
2.2 Initiative portfolio
● **Initiative:** Description [ref](URL)
● **Why it matters:** Operational impact [ref](URL)
● **KPI watchlist:** Metrics [ref](URL)
**Why this matters for Speridian**
Explain service opportunities [ref](URL)
STRICT RULES:
Every claim must include [ref](URL)
No hallucinations""",
"Section 3: Account Tech Landscape":"""Act as fintech architect.Use EXACT structure:
Section 3: Account Tech Landscape
Executive snapshot
Summary [ref](URL)
3.1 Digital banking
● **Online banking:** Details [ref](URL)
● **Mobile banking:** Details [ref](URL)
3.2 Core banking and partners
● **Core provider:** FIS,Fiserv,Jack Henry,etc [ref](URL)
● **Fintech partners:** Vendors [ref](URL)
3.3 Implications
Explain modernization gaps and opportunities [ref](URL)
STRICT RULES:
Every fact must include [ref](URL)
Use verified sources""",
"Section 4: Speridian Account Relationship":"""Act as sales director.Use EXACT structure:
Section 4: Speridian Account Relationship
4.1 Relationship analysis
Explain likely engagement areas [ref](URL)
4.2 Stakeholder roles
● CIO,CTO,Head of Digital Banking roles and priorities [ref](URL)
4.3 Strategic fit
Explain where Speridian can deliver value [ref](URL)
STRICT RULES:
DO NOT generate LinkedIn URLs
DO NOT generate personal contact info
Use role-level insights only
Include [ref](URL)""",
"Section 5: Speridian Next Steps":"""Act as managing director.Use EXACT structure:
Section 5: Speridian Next Steps
5.1 Funding signals
● Signal description [ref](URL)
● Why it creates opportunity [ref](URL)
5.2 Recommended plays
● Play description [ref](URL)
STRICT RULES:
Every claim must include [ref](URL)
No hallucinations"""}
def initializer(state:OverallState):
 return{"remaining_sections":list(PROMPT_SOP.keys()),"completed_research":[],"all_urls":[]}
def researcher_node(state:OverallState):
 if not state["remaining_sections"]:return state
 current_section=state["remaining_sections"][0]
 llm=get_llm()
 fintech_domains=["fdic.gov","sec.gov","americanbanker.com","finextra.com","bankingdive.com","fintechfutures.com","bankautomationnews.com"]
 if"Tech Landscape"in current_section:
  query=f"{state['target_company']} core banking provider FIS Fiserv Jack Henry fintech"
 elif"Business Overview"in current_section:
  query=f"{state['target_company']} annual report investor relations assets deposits net income"
 else:
  query=f"{state['target_company']} {current_section} banking strategy initiatives"
 search_tool=TavilySearchResults(max_results=5,include_domains=fintech_domains)
 try:
  results=search_tool.invoke({"query":query})
  urls=[];context=""
  for r in results:
   url=r.get("url","")
   urls.append(url)
   context+=f"\nSource [{url}]:{r.get('content','')}"
 except Exception as e:
  urls=[];context=f"Search failed:{str(e)}"
 sys_prompt=PROMPT_SOP[current_section]
 user_msg=f"Company:{state['target_company']}\nPDF Context:{state['pdf_context'][:4000]}\nVerified Sources:{context}"
 response=llm.invoke([("system",sys_prompt),("user",user_msg)])
 return{"completed_research":state["completed_research"]+[{"section":current_section,"content":response.content}],"all_urls":state["all_urls"]+urls,"remaining_sections":state["remaining_sections"][1:]}
def reflection_node(state:OverallState):
 llm=get_llm()
 latest=state["completed_research"][-1]
 audit_prompt=f"""You are audit analyst.Remove hallucinations.Remove claims without sources.Ensure every fact has [ref](URL).Return corrected version only:\n{latest['content']}"""
 response=llm.invoke([("user",audit_prompt)])
 state["completed_research"][-1]["content"]=response.content
 return state
def writer_node(state:OverallState):
 report=f"# Strategic Intelligence Report: {state['target_company']}\n"
 for item in state["completed_research"]:
  report+=f"\n## {item['section']}\n{item['content']}\n"
 report+="\n# References\n"
 for url in list(set(state["all_urls"])):report+=f"- {url}\n"
 return{"final_report":report}
def add_hyperlink(paragraph,url,text):
 part=paragraph.part
 r_id=part.relate_to(url,RELATIONSHIP_TYPE.HYPERLINK,is_external=True)
 hyperlink=OxmlElement("w:hyperlink")
 hyperlink.set(qn("r:id"),r_id)
 run=OxmlElement("w:r")
 text_el=OxmlElement("w:t")
 text_el.text=text
 run.append(text_el)
 hyperlink.append(run)
 paragraph._p.append(hyperlink)
def save_report_as_docx(final_text,target):
 from docx.shared import Pt,RGBColor,Inches
 from docx.enum.text import WD_ALIGN_PARAGRAPH
 from docx.oxml import OxmlElement
 from docx.enum.style import WD_STYLE_TYPE
 doc=Document()
 section=doc.sections[0]
 section.top_margin=Inches(1)
 section.bottom_margin=Inches(1)
 section.left_margin=Inches(1)
 section.right_margin=Inches(1)
 title=doc.add_paragraph()
 run=title.add_run("Strategic Intelligence Report")
 run.font.size=Pt(32)
 run.font.bold=True
 run.font.color.rgb=RGBColor(0,51,102)
 title.alignment=WD_ALIGN_PARAGRAPH.CENTER
 subtitle=doc.add_paragraph()
 run2=subtitle.add_run(target)
 run2.font.size=Pt(20)
 run2.font.color.rgb=RGBColor(89,89,89)
 subtitle.alignment=WD_ALIGN_PARAGRAPH.CENTER
 doc.add_page_break()
 lines=final_text.split("\n")
 for line in lines:
  line=line.strip()
  if not line:
   continue
  if line.startswith("# Strategic Intelligence Report"):
   continue
  if line.startswith("# References"):
   doc.add_page_break()
   ref=doc.add_paragraph()
   r=ref.add_run("References")
   r.font.size=Pt(18)
   r.font.bold=True
   r.font.color.rgb=RGBColor(0,51,102)
   continue
  if line.startswith("## "):
   h=doc.add_paragraph()
   r=h.add_run(line.replace("## ",""))
   r.font.size=Pt(18)
   r.font.bold=True
   r.font.color.rgb=RGBColor(0,51,102)
   continue
  if line.startswith("●"):
   p=doc.add_paragraph(style=None)
   run=p.add_run("● ")
   run.font.bold=True
   content=line[1:].strip()
  else:
   p=doc.add_paragraph()
   content=line
  parts=re.split(r'(\[ref\]\(.*?\))',content)
  for part in parts:
   link_match=re.match(r'\[ref\]\((.*?)\)',part)
   if link_match:
    url=link_match.group(1)
    part_rel=p.part
    r_id=part_rel.relate_to(url,RELATIONSHIP_TYPE.HYPERLINK,is_external=True)
    hyperlink=OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'),r_id)
    new_run=OxmlElement('w:r')
    rPr=OxmlElement('w:rPr')
    c=OxmlElement('w:color')
    c.set(qn('w:val'),'0563C1')
    rPr.append(c)
    u=OxmlElement('w:u')
    u.set(qn('w:val'),'single')
    rPr.append(u)
    new_run.append(rPr)
    text_el=OxmlElement('w:t')
    text_el.text=" [ref]"
    new_run.append(text_el)
    hyperlink.append(new_run)
    p._p.append(hyperlink)
   else:
    bold_parts=re.split(r'(\*\*.*?\*\*)',part)
    for bp in bold_parts:
     if bp.startswith("**") and bp.endswith("**"):
      run=p.add_run(bp.replace("**",""))
      run.bold=True
     else:
      run=p.add_run(bp)
      run.font.size=Pt(11)
 buf=io.BytesIO()
 doc.save(buf)
 return buf.getvalue()

workflow=StateGraph(OverallState)
workflow.add_node("initializer",initializer)
workflow.add_node("researcher",researcher_node)
workflow.add_node("reflection",reflection_node)
workflow.add_node("writer",writer_node)
workflow.set_entry_point("initializer")
workflow.add_edge("initializer","researcher")
workflow.add_edge("researcher","reflection")
workflow.add_conditional_edges("reflection",lambda x:"researcher"if x["remaining_sections"]else"writer",{"researcher":"researcher","writer":"writer"})
workflow.add_edge("writer",END)
app=workflow.compile()
st.markdown('<div class="container glass-card">',unsafe_allow_html=True)

col1,col2,col3=st.columns([2,2,1])

with col1:
 target=st.text_input("Target Company")

with col2:
 pdf=st.file_uploader("Upload Annual Report",type="pdf")

with col3:
 run=st.button("Run Analysis")

st.markdown('</div>',unsafe_allow_html=True)

if run and target:

 pdf_text=""

 if pdf:

  reader=PdfReader(pdf)

  pdf_text="\n".join([p.extract_text()for p in reader.pages if p.extract_text()])

 state={"target_company":target,"pdf_context":pdf_text,"remaining_sections":[],"completed_research":[],"all_urls":[],"final_report":""}

 final=""

 with st.status("Running strategic intelligence analysis...",expanded=True):
   

  for event in app.stream(state):

   for node,out in event.items():

    if node=="reflection":

     latest=out["completed_research"][-1]

     st.markdown(f"""

     <div class="container glass-card">

     <h4>{latest['section']}</h4>

     {latest['content']}

     </div>

     """,unsafe_allow_html=True)

    if node=="writer":

     final=out["final_report"]

 if final:

  st.download_button(

   "Download DOCX",

   save_report_as_docx(final,target),

   file_name=f"{target}_Strategic_Report.docx"

  )
